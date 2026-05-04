import pandas as pd
from docx import Document


# ─────────────────────────────────────────────
# DOCX PARSER (policy documents)
# ─────────────────────────────────────────────

def parse_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables cleanly
    for table in doc.tables:
        table_rows = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_rows.append(" | ".join(cells))

        if table_rows:
            parts.append("\n".join(table_rows))

    # Final clean join
    return "\n\n".join(parts).strip()


# ─────────────────────────────────────────────
# TXT PARSER
# ─────────────────────────────────────────────

def parse_txt(file_path: str) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue

    with open(file_path, "rb") as f:
        return f.read().decode("utf-8", errors="replace").strip()


# ─────────────────────────────────────────────
# XLSX (INTENTIONALLY DISABLED)
# ─────────────────────────────────────────────

def parse_xlsx(file_path: str):
    raise RuntimeError(
        "parse_xlsx is disabled. Use pandas-based retrieval (db.py) instead."
    )